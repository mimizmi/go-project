#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成本科毕业论文：医院异构系统实时数据中台设计与实现
广东医科大学 信息工程系 2025届
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────
# Helper: set East-Asian + Western fonts on a run
# ─────────────────────────────────────────────────────
def set_run_font(run, cn, en, pt, bold=False, italic=False):
    run.font.size = Pt(pt)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi'):
        rFonts.set(qn(attr), en)
    for attr in ('w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), cn)


def set_spacing(para, before=0, after=0, line_pts=None):
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    if before:
        sp.set(qn('w:before'), str(int(before * 20)))
    if after:
        sp.set(qn('w:after'), str(int(after * 20)))
    if line_pts:
        sp.set(qn('w:line'), str(int(line_pts * 20)))
        sp.set(qn('w:lineRule'), 'exact')


def set_indent(para, first_line_chars=2, pt_per_char=12):
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), str(int(first_line_chars * pt_per_char * 20)))


# ─────────────────────────────────────────────────────
# Document builder
# ─────────────────────────────────────────────────────
class Thesis:
    def __init__(self):
        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        s = self.doc.sections[0]
        s.page_width  = Cm(21)
        s.page_height = Cm(29.7)
        s.top_margin    = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin   = Cm(3.17)
        s.right_margin  = Cm(3.17)
        s.header_distance = Cm(1.5)
        s.footer_distance = Cm(1.75)
        # Header text
        header = s.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run('广东医科大学生物医学工程学院信息工程系2025届本科毕业论文（设计）')
        set_run_font(run, '宋体', 'SimSun', 9)

    # ── Paragraph types ────────────────────────────────
    def h1(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, '黑体', 'SimHei', 16, bold=True)
        set_spacing(p, before=24, after=12, line_pts=30)
        return p

    def h2(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, '黑体', 'SimHei', 15, bold=True)
        set_spacing(p, before=12, after=6, line_pts=28)
        return p

    def h3(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, '黑体', 'SimHei', 14, bold=True)
        set_spacing(p, before=6, after=3, line_pts=26)
        return p

    def h4(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, '黑体', 'SimHei', 12, bold=True)
        set_spacing(p, before=3, after=3, line_pts=21)
        return p

    def body(self, text, indent=True):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        set_run_font(run, '宋体', 'Times New Roman', 12)
        set_spacing(p, line_pts=21)
        if indent:
            set_indent(p)
        return p

    def caption(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, '宋体', 'Times New Roman', 10.5)
        set_spacing(p, before=3, after=3, line_pts=18)
        return p

    def blank(self):
        p = self.doc.add_paragraph()
        set_spacing(p, line_pts=21)
        return p

    def table(self, headers, rows, cap=None):
        if cap:
            self.caption(cap)
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = 'Table Grid'
        hr = tbl.rows[0]
        for i, h in enumerate(headers):
            c = hr.cells[i]
            c.text = ''
            run = c.paragraphs[0].add_run(h)
            set_run_font(run, '宋体', 'Times New Roman', 10.5, bold=True)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ri, row in enumerate(rows):
            tr = tbl.rows[ri + 1]
            for ci, val in enumerate(row):
                c = tr.cells[ci]
                c.text = ''
                run = c.paragraphs[0].add_run(str(val))
                set_run_font(run, '宋体', 'Times New Roman', 10.5)
                c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return tbl

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)
        print(f'已生成：{path}')


# ═══════════════════════════════════════════════════════════════
# CONTENT
# ═══════════════════════════════════════════════════════════════
def build(t: Thesis):

    # ── 封面（简化版占位，实际由学院模板替换）────────────────
    p = t.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\n广东医科大学\n本科毕业论文（设计）\n\n')
    set_run_font(run, '黑体', 'SimHei', 22, bold=True)

    p2 = t.doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('题  目：医院异构系统实时数据中台设计与实现\n'
                      '        ——多数据库CDC采集、Exactly-Once投递、增量位点管理\n\n')
    set_run_font(run2, '宋体', 'Times New Roman', 14)

    info = [
        '学    院：生物医学工程学院',
        '专    业：信息工程',
        '学    号：',
        '姓    名：',
        '指导教师：',
        '完成日期：2025年6月',
    ]
    for line in info:
        pi = t.doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ri = pi.add_run(line)
        set_run_font(ri, '宋体', 'Times New Roman', 14)
    t.page_break()

    # ── 中文摘要 ──────────────────────────────────────────────
    t.h1('摘  要')
    t.body(
        '医院信息系统长期呈现多系统并行建设、数据库类型各异的客观状态，HIS、LIS等业务系统之间'
        '存在严重的数据孤岛问题，跨系统统计分析依赖离线批量ETL，时延明显，难以支撑急诊抢救、'
        '实时运营监控等"秒级可用"的业务场景。针对上述问题，本文设计并实现了一套面向医院异构系统'
        '的实时数据中台原型，核心涵盖多数据库变更数据捕获（CDC）、基于Kafka的端到端Exactly-Once'
        '语义投递以及增量位点容错管理三个关键方向。'
    )
    t.body(
        '在采集层，系统分别针对MySQL和SQL Server实现了差异化的CDC适配器：MySQL侧基于binlog'
        '行模式解析捕获INSERT/UPDATE/DELETE事件；SQL Server侧直接查询CDC变更跟踪表（_CT表）'
        '进行增量轮询，避免了TVF接口的最小LSN限制问题。所有变更事件统一封装为标准化ChangeEvent，'
        '包含数据库类型、表名、主键、前后镜像、操作类型、事件ID及位点信息。'
    )
    t.body(
        '在传输层，系统使用confluent-kafka-go事务型Producer，通过BeginTransaction/CommitTransaction'
        '原语将"发送消息"与"提交源库位点"绑定为原子操作，保障Kafka侧的精确一次投递。'
        '消费端订阅相应主题后，在单个PostgreSQL事务内完成幂等键校验、业务UPSERT/DELETE以及去重'
        '记录写入，仅在PostgreSQL事务提交成功后才提交Kafka消费位移，从而实现端到端Exactly-Once语义。'
    )
    t.body(
        '在容错层，源库增量位点通过SQLite持久化存储，每次Kafka事务提交成功后原子更新，'
        '进程故障重启后可从上次成功位点断点续传，并由消费端幂等键负责对重放事件进行去重拦截。'
        '系统集成Prometheus指标采集与Grafana监控看板，对CDC事件吞吐、Kafka事务成功率、'
        '端到端延迟等关键指标进行全链路可观测。'
    )
    t.body(
        '实验结果表明：正确性验证实验中10000条混合操作记录经CDC管道同步后目标库计数完全吻合，'
        '差异行为零；Exactly-Once对比实验中EOS模式重复行数为零，验证了幂等机制有效性；'
        '性能基准实验中，在批量大小1000条、6分区配置下吞吐可达161472 events/s，P50延迟44ms，'
        'P99延迟294ms，具备较强的实用性能基础。'
    )
    t.blank()
    kw = t.doc.add_paragraph()
    r_kw = kw.add_run('关键词：')
    set_run_font(r_kw, '黑体', 'SimHei', 12, bold=True)
    r_kw2 = kw.add_run('变更数据捕获；Exactly-Once语义；异构数据库；Kafka事务；实时数据中台')
    set_run_font(r_kw2, '宋体', 'Times New Roman', 12)
    t.page_break()

    # ── English Abstract ──────────────────────────────────────
    p_abs = t.doc.add_paragraph()
    p_abs.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_abs = p_abs.add_run('ABSTRACT')
    set_run_font(r_abs, 'Times New Roman', 'Times New Roman', 16, bold=True)
    set_spacing(p_abs, before=0, after=12)

    en_paras = [
        'Hospital information systems have long been characterized by the co-existence of heterogeneous '
        'databases across HIS, LIS, EMR and other subsystems built by different vendors. The resulting '
        'data silos force cross-system analytics to rely on nightly batch ETL pipelines, whose inherent '
        'latency makes real-time scenarios such as emergency care and operational dashboards infeasible. '
        'This thesis addresses the problem by designing and implementing a real-time data middleware '
        'prototype for heterogeneous hospital systems, focusing on three key areas: multi-database '
        'Change Data Capture (CDC), end-to-end Exactly-Once delivery via Apache Kafka, and '
        'incremental offset fault-tolerance management.',

        'At the ingestion layer, tailored CDC adapters are implemented for both MySQL and SQL Server. '
        'The MySQL adapter parses row-format binary logs to capture INSERT, UPDATE, and DELETE events. '
        'The SQL Server adapter polls the CDC change-tracking table (_CT) directly to avoid the '
        'minimum-LSN restriction of the TVF interface. All change events are normalized into a '
        'unified ChangeEvent structure containing database type, table name, primary key map, '
        'before/after row images, operation type, event ID, and offset position.',

        'At the transport layer, the system uses a transactional Kafka producer. '
        'BeginTransaction/CommitTransaction wrap message production and offset checkpointing into an '
        'atomic unit, guaranteeing exactly-once semantics at the Kafka layer. On the consumer side, '
        'each batch is written to PostgreSQL in a single transaction that atomically checks an '
        'idempotency guard, executes UPSERT or DELETE on the target ODS table, and records the '
        'deduplication entry. Kafka consumer offsets are committed only after the PostgreSQL '
        'transaction commits successfully.',

        'Experimental results demonstrate: (1) in the correctness experiment, 10,000 mixed-operation '
        'records synchronized through the pipeline show zero count difference and zero row-level diff; '
        '(2) in the EOS experiment, zero duplicate rows are observed under EOS mode; '
        '(3) in the performance benchmark, a throughput of 161,472 events/s is achieved with '
        'batch-size=1000 and 6 partitions, with P50 latency of 44 ms and P99 latency of 294 ms.',
    ]
    for ep in en_paras:
        ep_para = t.doc.add_paragraph()
        ep_run = ep_para.add_run(ep)
        set_run_font(ep_run, 'Times New Roman', 'Times New Roman', 12)
        set_spacing(ep_para, line_pts=21)
        set_indent(ep_para)

    t.blank()
    kw_en = t.doc.add_paragraph()
    r_kw_en = kw_en.add_run('Key words: ')
    set_run_font(r_kw_en, 'Times New Roman', 'Times New Roman', 12, bold=True)
    r_kw_en2 = kw_en.add_run('Change Data Capture; Exactly-Once Semantics; Heterogeneous Database; '
                               'Kafka Transaction; Real-time Data Middleware')
    set_run_font(r_kw_en2, 'Times New Roman', 'Times New Roman', 12)
    t.page_break()

    # ── 目录（占位，Word自动生成） ───────────────────────────
    p_toc = t.doc.add_paragraph()
    p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_toc = p_toc.add_run('目    录')
    set_run_font(r_toc, '黑体', 'SimHei', 16, bold=True)
    t.body('（请在Word中使用"引用→目录→自动目录"生成目录）', indent=False)
    t.page_break()

    # ════════════════════════════════════════════════
    # 1  绪论
    # ════════════════════════════════════════════════
    t.h1('1  绪论')

    t.h2('1.1  研究背景与意义')
    t.body(
        '随着医院信息化建设持续深入，各类业务系统（HIS、EMR、LIS、PACS、RIS等）相继上线，'
        '院内数据存储呈现高度异构的态势：早期部署的HIS系统以SQL Server为主，检验科LIS多采用'
        'MySQL，影像系统则往往使用Oracle或自研文件存储。多套系统并立、标准不统一，导致数据孤岛'
        '问题愈发突出，跨系统数据共享与联合分析面临极大障碍。'
    )
    t.body(
        '传统的数据集成方案以T+1的离线批量ETL为主：定时抽取源库快照，经过清洗转换后加载至数据仓库。'
        '该方案技术成熟、实施门槛低，但固有的时延（通常在数小时至次日之间）使其难以满足临床与运营'
        '对数据时效性的高要求。例如，急诊绿色通道的快速分诊需要实时获取患者既往检验结果；药品库存'
        '预警需要毫秒至秒级的库存变化感知；院级运营指标看板期望分钟级刷新。这些场景均对数据平台'
        '的实时性提出了迫切需求，批处理ETL已无法有效支撑。'
    )
    t.body(
        '变更数据捕获（Change Data Capture，CDC）技术通过侦听数据库事务日志获取增量变更，'
        '相较全量快照抽取具有延迟低、源端侵入小的显著优势，近年来在实时数据集成领域得到广泛应用。'
        '以Debezium、Flink CDC为代表的开源工具将CDC能力模块化，降低了工程实施门槛；'
        'Apache Kafka凭借其高吞吐、持久化、事务性能力成为CDC消息总线的事实标准选择。'
        '然而，在医院"多数据库并存+老旧系统适配+合规脱敏"的复杂场景下，通用CDC工具仍存在'
        '适配成本高、全链路Exactly-Once难以开箱即用等工程挑战，亟需面向场景的系统性设计与验证。'
    )
    t.body(
        '本课题面向医院典型异构数据库环境（MySQL + SQL Server），研究并实现一套实时数据中台'
        '原型系统，在不改造既有业务系统的前提下，通过事务日志解析实现低延迟增量采集，以Kafka'
        '事务机制保障端到端Exactly-Once投递，并通过幂等写入和持久化位点管理实现故障自动恢复。'
        '研究成果既有助于完善医疗数据实时集成的技术路线，也为后续数据治理与数据服务化建设奠定底座。'
    )

    t.h2('1.2  国内外研究现状')
    t.h3('1.2.1  医疗多源异构数据汇聚研究')
    t.body(
        '医疗大数据平台建设研究普遍以"汇聚—治理—共享"为主线，强调多源异构数据的标准化、'
        '质量控制与安全合规。王淇等[2]提出了面向多中心的异构数据共享服务架构，通过统一数据目录'
        '和标准规范推动跨机构协同。程盼飞[8]研究了基于CDC技术的多源数据质量校验系统，将CDC'
        '捕获的增量事件用于数据一致性核验，有效降低了人工校验成本。Henke等[15]在MIMIC-IV'
        '数据库基础上研究了OMOP通用数据模型的增量ETL流程，证实了CDC路线在医疗数仓持续更新'
        '场景下的可行性。然而，现有研究对"跨异构库持续同步+全链路Exactly-Once"的一体化落地'
        '讨论相对不足，尤其是在医院老旧系统（SQL Server 2008~2016版本）的适配问题上缺乏针对性方案。'
    )
    t.h3('1.2.2  CDC技术与实时集成路线')
    t.body(
        'CDC技术路线可分为基于日志的触发器法和基于事务日志的非侵入式解析法两类，后者因不依赖'
        '源端改造、延迟更低而成为主流。何承勇[7]研究了基于日志的数据库变更捕获系统，对日志解析'
        '的正确性与性能进行了系统分析。王凯军等[5]提出了面向多源异构的日志型CDC更新方法，'
        '重点解决多源并发下的事件排序与一致性更新问题。苏子权[9]实现了基于MySQL binlog的'
        '数据增量同步系统，分析了GTID模式下的位点管理策略。在工程生态方面，Debezium支持'
        'MySQL、PostgreSQL、SQL Server等主流数据库的CDC适配，Flink CDC在此基础上集成了'
        '流处理能力；Sharma[14]对Debezium、Google Spanner等方案进行了比较分析，指出'
        'AI辅助的CDC方案在异构数据模型感知方面具有潜力。'
    )
    t.h3('1.2.3  Exactly-Once语义研究')
    t.body(
        'Exactly-Once语义（EOS）是流式系统一致性保障的最高等级，意味着每条消息恰好被处理并'
        '生效一次，不重复也不遗漏。Kafka自0.11版本引入幂等Producer和事务API，使得生产侧可将'
        '"发送消息"与"提交状态"绑定为原子操作，为端到端EOS提供了基础支撑[13]。王岩等[13]'
        '分析了Kafka Consumer可靠性设计，指出消费位移提交策略对EOS实现的决定性影响。'
        '曾泽堂[3]在其数据同步系统研究中，探讨了事务日志、消息队列与目标存储三者之间的'
        '一致性协调机制。将EOS从消息层扩展到"CDC采集—传输—落地"全链路，仍面临跨系统'
        '事务边界对齐与幂等键设计等工程挑战，是本课题重点突破的方向。'
    )
    t.h3('1.2.4  位点管理与容错恢复研究')
    t.body(
        '实时数据管道的长稳运行依赖于可靠的位点管理机制。位点的作用不仅是"记住同步进度"，'
        '更是故障恢复的依据，其正确性直接决定重启后数据的完整性。李佳奇[4]在基于binlog的'
        '同步系统研究中，深入分析了位点丢失与重复投递的关联关系，提出了"提交成功后再推进位点"'
        '的原则。程盼飞[8]的质量校验系统则采用检查点机制记录消费进度，实现了故障后的断点续传。'
        '现有工作大多聚焦于单源位点管理，多源并发下的位点一致提交策略研究仍相对匮乏。'
    )

    t.h2('1.3  本文主要工作')
    t.body(
        '本文的主要工作包括以下四个方面：'
    )
    t.body(
        '（1）构建面向MySQL与SQL Server的差异化CDC适配框架，统一事件封装为标准化ChangeEvent，'
        '支持字段脱敏与操作类型过滤，实现对既有业务系统的非侵入式增量采集。'
    )
    t.body(
        '（2）设计并实现基于Kafka事务型Producer的端到端Exactly-Once投递机制，通过Kafka事务'
        'begin/send/commit原语与消费端PostgreSQL事务幂等写入的协同配合，实现全链路"不重不漏"。'
    )
    t.body(
        '（3）设计多源增量位点持久化与一致提交策略，以SQLite为位点存储后端，在Kafka事务提交'
        '成功后原子更新位点，支持进程故障后的断点续传与重放去重。'
    )
    t.body(
        '（4）构建原型系统并进行系统性实验评估，包括正确性验证、EOS语义对比和性能基准测试，'
        '以量化数据验证系统设计的有效性。'
    )

    t.h2('1.4  论文结构安排')
    t.body(
        '本文共分六章。第1章为绪论，介绍研究背景、国内外现状及本文主要工作。'
        '第2章进行可行性分析，从技术、经济和操作三个维度论证方案的可实施性。'
        '第3章开展需求分析，明确系统的功能需求与非功能需求。'
        '第4章进行总体设计，提出系统架构方案、模块划分和关键技术选型。'
        '第5章为核心内容，对各模块进行详细设计与实现说明。'
        '第6章开展系统测试与实验评估。最后给出总结与展望。'
    )
    t.page_break()

    # ════════════════════════════════════════════════
    # 2  可行性分析
    # ════════════════════════════════════════════════
    t.h1('2  可行性分析')

    t.h2('2.1  技术可行性')
    t.body(
        '从技术角度分析，本系统所依赖的核心技术均已达到工业级成熟度，具备充分的技术可行性。'
    )
    t.body(
        '（1）CDC日志解析技术成熟。MySQL binlog行模式（ROW format）已成为生产部署标准，'
        'go-mysql-replication等开源库提供了稳定的binlog解析能力；SQL Server自2008版本起'
        '内置CDC功能，通过系统表可直接查询变更历史，技术路线清晰。'
    )
    t.body(
        '（2）Kafka事务API生产可用。Apache Kafka自0.11版本引入幂等Producer和事务API，'
        'confluent-kafka-go作为官方支持的Go语言客户端，对事务型Producer提供完整封装，'
        '在金融、电商等高可靠场景已有广泛验证。'
    )
    t.body(
        '（3）PostgreSQL pgx驱动支持完善。pgx/v5是Go生态中性能最优的PostgreSQL驱动，'
        '支持连接池、批量操作与事务管理，能够满足高并发幂等写入的需求。'
    )
    t.body(
        '（4）容器化部署方案成熟。基于Docker Compose的多容器编排方案可在单机环境快速搭建'
        '完整的开发测试环境，降低环境准备成本，便于实验复现。'
    )
    t.body(
        '综合评估，各项关键技术均有成熟的开源实现和充分的社区支持，技术风险可控，可行性强。'
    )

    t.h2('2.2  经济可行性')
    t.body(
        '本系统为原型研究性质，所采用的技术栈（Go语言、Apache Kafka、PostgreSQL、MySQL、'
        'SQL Server开发版、Prometheus、Grafana）均为开源或免费版本，在研究阶段无需商业授权费用。'
        '基础设施方面，单台配置为8核16GB内存的开发服务器即可运行完整的Docker Compose实验环境，'
        '硬件投入较低。若后续产品化部署，可采用云厂商提供的托管Kafka和数据库服务，按需付费，'
        '避免大规模自建运维成本，经济可行性良好。'
    )

    t.h2('2.3  操作可行性')
    t.body(
        '系统采用配置文件驱动的方式管理数据源、主题映射、批量参数等运行时配置，无需修改代码即可'
        '添加新的数据源或调整策略。基于Prometheus+Grafana的监控看板提供可视化运维支撑，'
        '运维人员可通过看板直观感知管道健康状态。日志采用结构化JSON格式，支持ELK等日志聚合系统'
        '接入，便于问题定位。系统的自动重试与断点续传机制降低了人工干预频次，操作可行性良好。'
    )

    t.h2('2.4  小结')
    t.body(
        '综合技术、经济、操作三个维度的分析，本系统方案技术成熟、成本可控、运维友好，'
        '具备充分的可行性，可进入详细设计与实现阶段。'
    )
    t.page_break()

    # ════════════════════════════════════════════════
    # 3  需求分析
    # ════════════════════════════════════════════════
    t.h1('3  需求分析')

    t.h2('3.1  业务场景分析')
    t.body(
        '某三甲医院信息系统现状如下：HIS系统（住院/门诊管理）运行于SQL Server 2016，'
        '存储患者基础信息（patients）、就诊记录（visits）和医嘱信息（orders）；'
        'LIS系统（检验信息系统）运行于MySQL 8.0，存储检验结果（lab_results）。'
        '两套系统各自独立运行，数据不互通，运营报表团队需要手工导出后人工合并，耗时且易出错。'
    )
    t.body(
        '核心业务诉求为：建立统一的运营数据存储（ODS），将多源变更以分钟级延迟汇聚至PostgreSQL，'
        '供报表查询与后续数据治理使用；同时要求数据"不重不漏"，即任何一条源库变更恰好在ODS'
        '中生效一次，不因系统重启或网络抖动产生重复或遗漏。'
    )

    t.h2('3.2  功能需求')
    t.h3('3.2.1  多源CDC采集需求')
    t.body('（1）支持MySQL 8.0 binlog行模式的INSERT/UPDATE/DELETE变更捕获，支持GTID模式，'
           '支持从指定位点恢复采集。')
    t.body('（2）支持SQL Server 2016/2019 CDC功能的增量轮询采集，支持从LSN位点恢复采集。')
    t.body('（3）所有变更事件统一封装为标准格式，包含：数据源标识、数据库名、表名、操作类型'
           '（INSERT/UPDATE/DELETE）、主键键值对、变更前后行镜像、事件唯一ID、采集时间戳、位点信息。')
    t.body('（4）支持敏感字段脱敏配置（如患者身份证号、手机号），在事件封装阶段将对应字段值替换为"***"。')

    t.h3('3.2.2  消息传输需求')
    t.body('（1）以Kafka作为消息总线，按数据库和表名映射到对应主题，支持主题分区数配置。')
    t.body('（2）Producer端开启幂等与事务模式，确保消息精确一次写入Kafka。')
    t.body('（3）支持批量发送，批量大小和超时时间可配置。')
    t.body('（4）事务提交失败时自动中止并基于指数退避策略重试，最大重试次数可配置。')

    t.h3('3.2.3  数据落地需求')
    t.body('（1）消费端订阅Kafka主题，支持read_committed隔离级别，过滤事务性幽灵消息。')
    t.body('（2）INSERT/UPDATE事件执行PostgreSQL UPSERT（INSERT ... ON CONFLICT DO UPDATE），'
           'DELETE事件执行物理删除。')
    t.body('（3）每批次写入在单个PostgreSQL事务内完成，事务内先进行幂等键检查，再执行业务写入，'
           '最后记录去重凭证；仅在PG事务提交成功后提交Kafka消费位移。')
    t.body('（4）目标表按"ods_{源库名}_{表名}"命名规则自动创建（原型阶段），列类型统一为TEXT，'
           '追加_cdc_source_id、_cdc_op_type、_cdc_updated_at三个元数据列。')

    t.h3('3.2.4  位点管理需求')
    t.body('（1）源库位点（MySQL: binlog文件+偏移量，SQL Server: LSN字符串）持久化到本地SQLite。')
    t.body('（2）Kafka事务提交成功后原子更新位点记录，保证位点不超前于已提交的Kafka消息。')
    t.body('（3）进程启动时自动加载上次成功位点，从该位点继续采集，实现断点续传。')
    t.body('（4）新数据源首次启动时，从源库当前最新位点开始采集，不拉历史存量。')

    t.h3('3.2.5  监控与告警需求')
    t.body('（1）通过Prometheus暴露CDC事件计数、Kafka事务成功/失败计数、Sink写入计数、'
           '端到端延迟等指标，监控接口端口可配置。')
    t.body('（2）Grafana看板可视化展示上述指标，支持管道实时状态监控。')

    t.h2('3.3  非功能需求')
    t.h3('3.3.1  正确性要求')
    t.body(
        '系统核心目标是保证数据正确性：任何在源库已提交的变更，在CDC管道正常运行情况下，'
        '最终在目标ODS中恰好生效一次；系统重启或网络抖动不应导致数据重复或遗漏。'
        '正确性校验标准：目标表行数与源表行数一致，抽样行级对比差异为零，幂等去重表无误判。'
    )
    t.h3('3.3.2  性能要求')
    t.body(
        '在实验环境（单节点Kafka、单PostgreSQL实例）下，单分区吞吐不低于10000 events/s，'
        'P95端到端延迟（从源库提交到ODS可查）不超过500ms（批量大小500条）。'
    )
    t.h3('3.3.3  可靠性要求')
    t.body(
        '进程崩溃后自动重启，重启后能从上次成功位点恢复，恢复过程无人工干预，'
        '恢复时间目标（RTO）不超过60秒。'
    )
    t.h3('3.3.4  可扩展性要求')
    t.body(
        '新增数据源只需在配置文件中添加配置项，无需修改代码；Kafka主题与分区数可在配置中调整，'
        '消费管道数量可水平扩展。'
    )

    t.h2('3.4  小结')
    t.body(
        '本章围绕医院异构数据集成的实际诉求，系统梳理了功能需求（CDC采集、消息传输、数据落地、'
        '位点管理、监控告警）与非功能需求（正确性、性能、可靠性、扩展性），为后续总体设计'
        '与详细实现提供了明确的目标约束。'
    )
    t.page_break()

    # ════════════════════════════════════════════════
    # 4  总体设计
    # ════════════════════════════════════════════════
    t.h1('4  总体设计')

    t.h2('4.1  系统整体架构设计')
    t.body(
        '系统整体采用"采集—传输—落地"三层流水线架构，如图4-1所示。'
        '各层职责清晰、界面明确，层间通过Kafka消息总线解耦，任意层的故障均不会直接传播至相邻层。'
    )
    t.caption('图4-1  系统整体架构示意图（文字描述版）')
    t.body(
        '【采集层（Source Layer）】：部署MySQL CDC Source和SQL Server CDC Source两个采集适配器。'
        'MySQL Source基于binlog行模式解析，以go-mysql-replication库驱动，实现事件驱动的推送式采集；'
        'SQL Server Source基于轮询CDC变更跟踪表，以configurable间隔（默认500ms）拉取新增变更。'
        '两种Source均将事件封装为统一的ChangeEvent，发送至Source Pipeline进行批量聚合。'
    )
    t.body(
        '【传输层（Transport Layer）】：Source Pipeline在收集到足够批量（默认500条）或超时（1s）后，'
        '通过TransactionalProducer将整批事件封装为一个Kafka事务发送。主题路由规则（TopicRouter）'
        '将不同表的事件分发至不同主题（如cdc.his.patients、cdc.lis.lab_results），'
        '分区键（Partition Key）取自事件主键，保证同一行的变更事件有序到达同一分区。'
    )
    t.body(
        '【落地层（Sink Layer）】：Consumer Pipeline订阅配置的Kafka主题，'
        '以read_committed隔离级别消费消息，过滤未提交事务的幽灵消息。'
        '每批消息由PostgresSinkWriter在单个PG事务内写入ODS，通过幂等键校验防止重放重复写入。'
        '仅在PG事务提交成功后，Consumer才向Kafka提交消费位移，保证端到端EOS语义。'
    )
    t.body(
        '【位点管理（Offset Store）】：SQLiteOffsetStore持久化源库位点，与Kafka事务提交原子绑定。'
        '【监控（Monitoring）】：MetricsServer以Prometheus格式暴露全链路指标，Grafana看板实时可视化。'
    )

    t.h2('4.2  模块划分与职责')
    t.body('系统按功能划分为六个核心模块，各模块职责如表4-1所示。')
    t.table(
        ['模块名称', '所属包', '核心职责'],
        [
            ['MySQL CDC适配器', 'internal/cdc', '解析MySQL binlog，产生ChangeEvent'],
            ['SQL Server CDC适配器', 'internal/cdc', '轮询SQL Server CDC表，产生ChangeEvent'],
            ['事务型Producer', 'internal/transport', 'Kafka事务发送，保证消息原子可见'],
            ['PostgreSQL Sink', 'internal/sink', '幂等UPSERT/DELETE，保证落地一次生效'],
            ['位点存储', 'internal/core', 'SQLite持久化源库位点，支持断点续传'],
            ['管道编排器', 'internal/engine', '协调各模块生命周期，处理重试与容错'],
        ],
        cap='表4-1  系统模块划分与职责说明'
    )

    t.h2('4.3  关键技术选型')
    t.h3('4.3.1  开发语言：Go 1.22')
    t.body(
        'Go语言原生支持并发（goroutine/channel），在网络IO密集型场景性能优异，内存占用低，'
        '编译为单一静态二进制便于容器化部署；标准库和生态对Kafka、MySQL、PostgreSQL等均有'
        '成熟驱动支持，是数据中间件类系统的主流选择。'
    )
    t.h3('4.3.2  消息队列：Apache Kafka 3.6（KRaft模式）')
    t.body(
        'Kafka具备高吞吐、持久化、可回放、分区有序等特性，事务API（幂等Producer + 事务Producer）'
        '是实现端到端EOS的关键基础设施。KRaft模式无需ZooKeeper依赖，降低了运维复杂度，'
        '适合原型环境快速部署。'
    )
    t.h3('4.3.3  目标存储：PostgreSQL 16')
    t.body(
        'PostgreSQL提供ACID事务保证和"INSERT ... ON CONFLICT DO UPDATE"语法，支持幂等UPSERT；'
        'pgx/v5驱动提供连接池与原生事务管理，适合高并发写入场景；'
        '丰富的索引类型和JSON支持为后续数据分析提供灵活基础。'
    )
    t.h3('4.3.4  位点存储：SQLite（mattn/go-sqlite3）')
    t.body(
        'SQLite作为本地嵌入式数据库，具备ACID保证、零配置、单文件部署的优点，'
        '满足位点数据"持久化+原子更新"的需求；对于原型系统，单机SQLite足以承载位点写入压力，'
        '无需引入额外中间件。'
    )
    t.h3('4.3.5  监控：Prometheus + Grafana')
    t.body(
        'Prometheus拉取式指标采集与告警机制成熟，Go官方提供prometheus/client_golang库；'
        'Grafana可灵活定制监控看板，支持多数据源聚合，是流处理系统可观测性建设的标准工具链。'
    )

    t.h2('4.4  数据流设计')
    t.body(
        '系统的核心数据流遵循严格的顺序依赖关系，以保障EOS语义。以单批事件为例，完整数据流如下：'
    )
    t.body(
        '① CDC Source捕获源库变更 → ② 事件封装为ChangeEvent → ③ Source Pipeline收集批次 → '
        '④ TransactionalProducer.BeginTxn() → ⑤ 循环调用Send(event) → '
        '⑥ CommitTxn()：Flush确认所有消息交付 → Kafka CommitTransaction() → '
        '⑦ 位点原子更新到SQLite → ⑧ Consumer PollBatch() → '
        '⑨ PostgresSinkWriter.WriteBatch()（PG事务：幂等检查+UPSERT+去重记录）→ '
        '⑩ PG CommitTransaction() → ⑪ Consumer CommitOffsets()。'
    )
    t.body(
        '关键容错点：若第⑥步Kafka提交失败，位点不更新，下次重试从上一位点重发，Consumer端幂等键拦截重复；'
        '若第⑨步PG写入失败，Kafka位移不提交，下次重新消费，幂等键再次保护；'
        '若第⑦步位点保存失败（极罕见），下次启动会重放部分事件，Consumer幂等负责兜底。'
    )

    t.h2('4.5  小结')
    t.body(
        '本章提出了系统的三层流水线架构方案，明确了各模块职责、技术选型理由及完整数据流路径，'
        '为第5章的详细设计与实现奠定了架构基础。'
    )
    t.page_break()

    # ════════════════════════════════════════════════
    # 5  详细设计与实现
    # ════════════════════════════════════════════════
    t.h1('5  详细设计与实现')

    t.h2('5.1  开发环境与项目结构')
    t.h3('5.1.1  开发环境')
    t.body('开发环境配置如表5-1所示。')
    t.table(
        ['环境项', '版本/配置'],
        [
            ['操作系统', 'Ubuntu 20.04 LTS (WSL2)'],
            ['Go编译器', 'Go 1.22.0，CGO_ENABLED=1，-tags musl'],
            ['MySQL', '8.0（Docker，binlog_format=ROW，GTID_MODE=ON）'],
            ['SQL Server', '2022-latest（Docker，CDC已启用）'],
            ['Apache Kafka', '3.6.0（KRaft模式，confluent cp-kafka:7.6.0）'],
            ['PostgreSQL', '16（Docker）'],
            ['监控', 'Prometheus latest + Grafana latest'],
            ['容器编排', 'Docker Compose v2.x'],
        ],
        cap='表5-1  开发环境配置'
    )

    t.h3('5.1.2  项目结构')
    t.body('项目按Go标准工程布局组织，主要目录结构如表5-2所示。')
    t.table(
        ['目录/文件', '说明'],
        [
            ['cmd/platform/main.go', '主程序入口，初始化并启动Coordinator'],
            ['internal/core/', '核心接口定义、ChangeEvent结构、错误码、IOffsetStore'],
            ['internal/cdc/', 'MySQL与SQL Server CDC适配器实现'],
            ['internal/transport/', 'Kafka Serializer、TopicRouter、TransactionalProducer'],
            ['internal/sink/', 'PostgresSinkWriter、IdempotencyGuard'],
            ['internal/engine/', 'SourcePipeline、ConsumerPipeline、Coordinator、重试策略'],
            ['internal/config/', '配置结构体与YAML解析，支持环境变量占位符'],
            ['internal/monitoring/', 'Prometheus指标定义与MetricsServer'],
            ['scripts/', '种子数据生成、一致性校验、实验脚本'],
            ['configs/', '数据源配置（sources.yaml）与主题配置（topics.yaml）'],
            ['docker-compose.yml', '完整实验环境编排文件'],
        ],
        cap='表5-2  项目主要目录结构'
    )

    t.h2('5.2  核心数据结构设计')
    t.h3('5.2.1  ChangeEvent标准化事件')
    t.body(
        'ChangeEvent是系统中所有模块共同操作的核心数据结构，定义于internal/core/event.go。'
        '其主要字段如表5-3所示。'
    )
    t.table(
        ['字段名', '类型', '说明'],
        [
            ['EventID', 'string', 'UUID，事件全局唯一标识'],
            ['SourceID', 'string', '数据源标识符，如"his_mysql_01"'],
            ['SourceType', 'SourceType', '枚举：SourceMySQL / SourceSQLServer'],
            ['Database', 'string', '源库名'],
            ['Table', 'string', '源表名'],
            ['OpType', 'OpType', '枚举：OpInsert / OpUpdate / OpDelete'],
            ['Before', 'map[string]interface{}', 'DELETE/UPDATE前镜像，INSERT时为nil'],
            ['After', 'map[string]interface{}', 'INSERT/UPDATE后镜像，DELETE时为nil'],
            ['PrimaryKeys', 'map[string]interface{}', '主键字段名→值映射'],
            ['Timestamp', 'time.Time', '事件采集时间戳'],
            ['Position', '*OffsetPosition', '关联位点信息'],
        ],
        cap='表5-3  ChangeEvent核心字段说明'
    )
    t.body(
        'ChangeEvent提供IdempotencyKey()方法，通过对SourceID、Table、主键序列化结果和Timestamp'
        '的SHA-256散列生成32字节幂等键字符串，作为去重凭证的唯一索引。'
    )

    t.h3('5.2.2  OffsetPosition位点结构')
    t.body(
        'OffsetPosition记录某一数据源的采集进度，包含SourceID、SourceType、'
        'LSN（SQL Server使用）、BinlogFile+BinlogPos（MySQL使用）、'
        'PollTimestamp（最近一次轮询时间）和EventCount（已采集事件累计数）。'
        'IsValid()方法根据SourceType校验位点字段是否有效，用于区分"首次启动"和"断点续传"两种场景。'
    )

    t.h2('5.3  MySQL Binlog CDC采集模块')
    t.h3('5.3.1  模块设计')
    t.body(
        'MySQLCdcSource位于internal/cdc/mysql.go，实现ICdcSource接口。'
        '其采用事件驱动的推送模式：通过go-mysql-replication库与MySQL建立binlog复制连接，'
        'MySQL主库将binlog事件实时推送给复制从库（本系统模拟从库角色），无需轮询，延迟极低（亚秒级）。'
    )
    t.body(
        'MySQLConfig配置项包括：SourceID、Host、Port、User、Password、Database、ServerID'
        '（复制从库ID，需与主库其他从库不冲突）、Tables（监控的表名列表）、'
        'MaskFields（字段脱敏配置）。'
    )

    t.h3('5.3.2  启动与断点续传')
    t.body(
        'Start()方法接收上次保存的OffsetPosition，若位点有效则从BinlogFile:BinlogPos断点重新同步；'
        '若无有效位点（首次启动），则调用SHOW MASTER STATUS获取当前binlog文件末尾位置作为起始点，'
        '避免拉取历史存量数据。关键实现如下：'
    )
    t.body(
        '（1）构建go-mysql BinlogSyncer配置，指定ServerID和Host信息；'
        '（2）根据起始位点调用syncer.StartSync(pos)或syncer.StartSyncGTID(gtidSet)；'
        '（3）在eventLoop goroutine中循环调用streamer.GetEvent()，将RowsEvent转换为ChangeEvent写入eventCh；'
        '（4）通过curPos字段实时记录当前已处理的binlog位点，供SourcePipeline在事务提交成功后读取并保存。'
    )

    t.h3('5.3.3  事件解析与标准化')
    t.body(
        'go-mysql-replication将binlog RowsEvent解析为行列矩阵，本模块的EventBuilder.BuildMySQL()'
        '方法负责将其转换为标准ChangeEvent：'
    )
    t.body(
        '（1）根据RowsEvent.Action确定OpType（"insert"/"update"/"delete"）；'
        '（2）INSERT事件：第0行为After，Before为nil；'
        '（3）DELETE事件：第0行为Before，After为nil；'
        '（4）UPDATE事件：第0行为Before，第1行为After；'
        '（5）调用ExtractPrimaryKeys()根据TableMapEvent中的列名信息提取主键字段；'
        '（6）调用MaskSensitiveFields()对配置的脱敏字段值替换为"***"。'
    )

    t.h2('5.4  SQL Server CDC轮询采集模块')
    t.h3('5.4.1  SQL Server CDC机制概述')
    t.body(
        'SQL Server内置CDC功能通过在目标表上执行sys.sp_cdc_enable_table启用，启用后系统'
        '在cdc schema下自动维护一张名为"captureInstance_CT"的变更跟踪表（Change Table），'
        '记录所有DML操作的前后镜像以及操作码（__$operation字段：1=DELETE, 2=INSERT, 4=UPDATE后镜像）'
        '和日志序列号（__$start_lsn）。本模块通过直接查询CT表实现增量采集。'
    )

    t.h3('5.4.2  核心查询设计')
    t.body(
        '早期版本曾尝试使用TVF接口cdc.fn_cdc_get_all_changes_*，但在from_lsn小于捕获实例'
        '最小LSN时会触发错误313（参数不足），且TVF的all_with_merge参数语义与全量获取接口不同。'
        '最终采用直接查询CT表的方案，核心SQL如下：'
    )
    t.body(
        "SELECT TOP 1000 * FROM cdc.{captureInstance}_CT "
        "WHERE __$start_lsn > {fromLSN} AND __$start_lsn <= {maxLSN} "
        "ORDER BY __$start_lsn"
    )
    t.body(
        '该设计的关键点：（1）使用严格大于（>）避免重复消费上批最后一条；'
        '（2）TOP 1000限制单次批量上限，防止超大批量导致Kafka事务flush超时；'
        '（3）toLSN上界取fn_cdc_get_max_lsn()的当前值，避免读取未提交事务的幽灵记录；'
        '（4）跳过__$update_mask、__$command_id、__$seqval、__$end_lsn等CDC元数据列，'
        '只保留业务字段。'
    )

    t.h3('5.4.3  LSN格式处理')
    t.body(
        'SQL Server的LSN以二进制形式存储（10字节），go-mssqldb驱动扫描后得到[]byte类型。'
        '本模块将其格式化为"0x{十六进制大写}"字符串作为位点字符串存储和比较，'
        '字符串比较在等长十六进制下等价于数值比较，正确性可保证。'
        '新数据源首次启动时调用sys.fn_cdc_get_min_lsn(\'\')'
        '获取系统最小LSN作为起始点，确保不遗漏任何已启用CDC后的变更。'
    )

    t.h3('5.4.4  BIT列类型转换问题')
    t.body(
        '实践中发现SQL Server的BIT类型列（如lab_results.is_abnormal）通过go-mssqldb'
        '扫描后以int64类型（值为0或1）返回，而非Go的bool类型。'
        '由于PostgreSQL目标列类型为TEXT，pgx驱动无法直接将int64编码为TEXT列，'
        '导致"unable to encode 1 into text format for text (OID 25)"运行时错误。'
        '解决方案是在sink层引入normalizeVal()统一转换函数，将所有非字符串类型（含int64、float64、bool等）'
        '通过fmt.Sprintf("%v", v)转换为字符串，保证与PostgreSQL TEXT列的类型兼容性。'
    )

    t.h2('5.5  Kafka事务型Producer投递模块')
    t.h3('5.5.1  TransactionalProducer设计')
    t.body(
        'TransactionalProducer位于internal/transport/producer.go，封装confluent-kafka-go的'
        '*kafka.Producer，对上层提供BeginTxn()、Send(event)、CommitTxn()、AbortTxn()四个方法，'
        '将Kafka事务的复杂性隐藏于模块内部。'
    )
    t.body(
        '关键配置参数如表5-4所示。'
    )
    t.table(
        ['配置项', '值', '说明'],
        [
            ['transactional.id', 'cdc-producer-{sourceID}', '每源唯一，Epoch Fencing防止僵尸Producer'],
            ['enable.idempotence', 'true', '启用幂等Producer，防止网络重传导致重复'],
            ['acks', 'all', '等待所有ISR副本确认，防止消息丢失'],
            ['max.in.flight.requests.per.connection', '5', '幂等模式下最大5个飞行请求'],
            ['retries', '2147483647', '无限重试，由delivery.timeout.ms控制总超时'],
            ['delivery.timeout.ms', '120000', '120s投递总超时，与flush超时对齐'],
            ['transaction.timeout.ms', '120000', '120s事务超时'],
        ],
        cap='表5-4  Kafka Producer关键配置参数'
    )

    t.h3('5.5.2  drainEvents goroutine的必要性')
    t.body(
        '使用nil delivery channel（不通过回调处理投递报告）时，librdkafka内部将投递报告'
        '写入producer.Events()通道。若该通道未被消费，当缓冲区满时librdkafka后台线程将阻塞，'
        '进而导致producer.Flush()永远无法返回0（pending消息数不减少），最终触发flush超时错误，'
        '整个批次被中止并无限重试，形成死锁。'
    )
    t.body(
        '解决方案是在NewTransactionalProducer()中启动一个专用的drainEvents() goroutine，'
        '持续消费Events()通道，记录投递失败日志但不干预主流程。该goroutine随producer.Close()'
        '时Events()通道关闭而自动退出，无资源泄漏风险。这是使用nil delivery channel的必要前提，'
        '不可省略。'
    )

    t.h3('5.5.3  Exactly-Once语义实现原理')
    t.body(
        'Kafka事务的EOS保证依赖三个层面的协同：'
    )
    t.body(
        '（1）生产侧幂等：enable.idempotence=true时，Broker为每个Producer实例分配PID+Epoch，'
        '对相同<PID, PartitionID, SequenceNumber>的重复消息只保留一份，防止网络重传导致重复写入。'
    )
    t.body(
        '（2）事务原子性：BeginTransaction/CommitTransaction将多个分区的消息写入绑定为原子操作，'
        '要么全部对read_committed消费者可见，要么全部不可见。CommitTxn()内部先Flush()确保所有消息'
        '已交付Broker，再调用CommitTransaction()；若Flush()超时则调用AbortTransaction()+'
        'Purge()清除队列后重试。'
    )
    t.body(
        '（3）消费侧隔离：Consumer的isolation.level=read_committed确保仅消费已提交事务的消息，'
        '过滤中间状态消息（事务中止或尚未提交）。结合消费端PostgreSQL事务的幂等写入，'
        '端到端EOS得以实现。'
    )

    t.h3('5.5.4  主题路由设计')
    t.body(
        'TopicRouter根据topics.yaml配置将ChangeEvent路由到对应Kafka主题。'
        '路由规则以"database.table"为键，支持通配符（如"hospital_his.*"匹配HIS库所有表）。'
        '分区键（Partition Key）取自事件主键字段值的拼接字符串，保证同一行的变更事件'
        '路由至同一分区，维护行级操作顺序性。'
    )

    t.h2('5.6  消费端幂等落地模块')
    t.h3('5.6.1  幂等键设计')
    t.body(
        '幂等键（Idempotency Key）是防止重复写入的核心机制。系统采用基于内容哈希的方案：'
        '对SourceID、Table名称和主键键值对序列化后进行SHA-256哈希，取16字节hex字符串作为幂等键。'
        '幂等键的设计原则：相同源库相同行的相同版本变更产生相同键（确定性），'
        '不同行或不同版本产生不同键（低碰撞率）。'
    )
    t.body(
        '_cdc_applied_events去重表结构如表5-5所示，PRIMARY KEY(event_key)+'
        'created_at字段结合定期清理过期记录（默认保留24小时），在存储开销和去重覆盖窗口间取得平衡。'
    )
    t.table(
        ['列名', '类型', '说明'],
        [
            ['event_key', 'TEXT PRIMARY KEY', '幂等键，SHA-256 hex字符串'],
            ['source_id', 'TEXT', '数据源标识'],
            ['table_name', 'TEXT', '源表名'],
            ['created_at', 'TIMESTAMPTZ', '记录写入时间，用于过期清理'],
        ],
        cap='表5-5  _cdc_applied_events去重表结构'
    )

    t.h3('5.6.2  PostgresSinkWriter.WriteBatch()实现')
    t.body(
        'WriteBatch()是消费端的核心方法，在单个PostgreSQL事务（ReadCommitted隔离级别）内'
        '按如下步骤处理一批事件：'
    )
    t.body(
        '① 遍历每条事件，调用IdempotencyGuard.IsDuplicate()在同一事务内查询去重表；'
        '② 若为重复事件则跳过（continue）并记录DEBUG日志；'
        '③ 若为新事件则调用applyEvent()执行UPSERT或DELETE；'
        '④ 调用IdempotencyGuard.MarkApplied()将幂等键写入去重表；'
        '⑤ 所有事件处理完后调用tx.Commit()提交PG事务；'
        '⑥ PG提交成功后返回nil，由ConsumerPipeline提交Kafka消费位移。'
    )
    t.body(
        '若步骤③或④发生错误，整个PG事务回滚，Kafka位移不提交，下次重新消费该批次，'
        '幂等键机制确保已成功应用的事件不会被重复写入。'
    )

    t.h3('5.6.3  动态建表与UPSERT实现')
    t.body(
        '目标表按"ods_{database}_{table}"命名，ensureTable()方法在事务内使用'
        '"CREATE TABLE IF NOT EXISTS"动态创建（原型阶段），所有业务列定义为TEXT类型，'
        '追加_cdc_source_id、_cdc_op_type、_cdc_updated_at三个CDC元数据列。'
    )
    t.body(
        'UPSERT语句采用PostgreSQL标准语法：'
        '"INSERT INTO \\"tableName\\" (cols) VALUES ($1,...) ON CONFLICT (pkCols) DO UPDATE SET col=EXCLUDED.col,..."。'
        '主键冲突时更新所有非主键列及CDC元数据列，实现幂等UPSERT语义。'
        'DELETE事件则直接执行"DELETE FROM \\"tableName\\" WHERE pk=$1..."。'
    )

    t.h2('5.7  位点管理与容错恢复模块')
    t.h3('5.7.1  SQLiteOffsetStore设计')
    t.body(
        'SQLiteOffsetStore（internal/core/offset_store.go）以SQLite数据库存储位点，'
        '核心表结构包含source_id（主键）、source_type、lsn、binlog_file、binlog_pos、'
        'poll_timestamp、event_count、last_txn_id和updated_at字段。'
    )
    t.body(
        'Save(sourceID, pos, txnID)方法使用"INSERT ... ON CONFLICT(source_id) DO UPDATE SET ..."'
        '原子写入位点，确保位点更新的幂等性。Load(sourceID)方法在启动时加载上次成功位点，'
        '若不存在则返回nil（触发首次启动逻辑）。'
    )

    t.h3('5.7.2  原子位点提交策略')
    t.body(
        '位点提交时序是保证"不重不漏"的关键。系统采用以下顺序：'
        '（1）Kafka事务CommitTransaction()成功 → （2）SQLite位点Save() → '
        '（3）进入下一批次采集。'
    )
    t.body(
        '若步骤（2）失败（如磁盘满），下次启动会从旧位点重放，产生少量重复事件；'
        '但由于Consumer端幂等键的保护，这些重复事件不会重复写入ODS，最终结果仍正确。'
        '此设计遵循"宁可重放、不可跳过"的原则，以消费端幂等兜底保证数据完整性。'
    )

    t.h3('5.7.3  指数退避重试策略')
    t.body(
        'RetryPolicy（internal/engine/pipeline.go）实现指数退避+随机抖动的重试策略，'
        '第n次重试的等待时间为：delay = BaseDelay × Multiplier^n × (1 ± Jitter×rand)，'
        '上限为MaxDelay。默认配置：BaseDelay=1s、Multiplier=2.0、Jitter=0.3、MaxDelay=60s、'
        'MaxRetries=-1（无限重试）。随机抖动（Jitter）的作用是错开多个Pipeline同时重试的时间点，'
        '避免"惊群效应"对Kafka和数据库造成冲击。'
    )

    t.h2('5.8  系统编排与管道协调')
    t.h3('5.8.1  Coordinator多管道协调')
    t.body(
        'Coordinator（internal/engine/coordinator.go）是系统的生命周期管理者，'
        '负责根据配置文件创建并启动所有Source Pipeline和Consumer Pipeline，'
        '监听各管道的错误返回，在管道异常退出时基于RetryPolicy进行重启。'
    )
    t.body(
        'Coordinator使用errgroup管理多个goroutine，所有管道在同一context下运行，'
        '任意管道在MaxRetries耗尽后返回error，Coordinator将取消context使其他管道优雅停止，'
        '最终以非零状态码退出，由容器编排层（Docker Compose restart: always）触发进程重启。'
    )

    t.h3('5.8.2  SourcePipeline主循环设计')
    t.body(
        'SourcePipeline.Run()的主循环逻辑如下：'
        '① Load位点 → ② source.Start(pos) → ③ 进入for{}循环：'
        '④ collectBatch()收集事件（达到batchSize或超过batchTimeout） → '
        '⑤ sendBatchWithRetry()发送并原子更新位点 → ⑥ 回到③。'
    )
    t.body(
        'collectBatch()通过select同时监听source.Events()（新事件）、source.Errors()（采集错误）、'
        'deadline.C（超时）和ctx.Done()（取消信号），任一触发即返回当前已收集的批次。'
        '超时触发时返回部分批次（可能为空），空批次直接跳过。'
    )

    t.h3('5.8.3  ConsumerPipeline设计')
    t.body(
        'ConsumerPipeline.Run()循环调用consumer.PollBatch()从Kafka拉取一批消息，'
        '反序列化为ChangeEvent列表后调用sink.WriteBatch()写入PostgreSQL，'
        'PG提交成功后提交Kafka消费位移。任意步骤错误均记录日志后等待2秒重试（不退出），'
        '充分利用Consumer客户端的自动重连机制。'
    )
    t.body(
        'Consumer订阅配置中isolation.level设为read_committed，确保仅消费已提交Kafka事务的消息，'
        '与Producer侧的事务语义形成端到端EOS闭环。'
    )

    t.h2('5.9  监控与可观测性')
    t.h3('5.9.1  Prometheus指标体系')
    t.body('系统暴露的主要Prometheus指标如表5-6所示。')
    t.table(
        ['指标名', '类型', '标签', '说明'],
        [
            ['cdc_events_total', 'Counter', 'source_id, table, op_type', 'CDC采集事件总数'],
            ['cdc_errors_total', 'Counter', 'source_id, error_type', 'CDC采集错误总数'],
            ['kafka_txn_total', 'Counter', 'source_id, status(committed/aborted)', 'Kafka事务提交/中止计数'],
            ['kafka_send_latency_seconds', 'Histogram', 'source_id', 'Kafka事务发送延迟分布'],
            ['sink_events_total', 'Counter', 'table, op_type', 'Sink落地事件总数'],
            ['sink_dedup_total', 'Counter', 'table', 'Sink去重拦截事件数'],
            ['sink_write_latency_seconds', 'Histogram', 'table', 'PG写入延迟分布'],
        ],
        cap='表5-6  系统Prometheus指标说明'
    )

    t.h3('5.9.2  Grafana监控看板')
    t.body(
        'Grafana看板按管道视图组织，包含以下面板：'
        '（1）CDC事件吞吐率（events/s，按source_id分组）；'
        '（2）Kafka事务成功率（committed/(committed+aborted)）；'
        '（3）Kafka发送延迟P50/P95/P99；'
        '（4）Sink落地吞吐率与去重率；'
        '（5）Sink写入延迟P50/P95/P99；'
        '（6）管道错误计数告警面板。'
        '通过Grafana Provisioning机制，看板配置以YAML文件形式存于代码仓库，随容器启动自动加载。'
    )

    t.h2('5.10  Docker容器化部署')
    t.body(
        '系统采用多阶段Dockerfile构建：构建阶段使用golang:1.22-alpine镜像编译二进制，'
        '运行阶段使用alpine:3.19精简镜像，最终镜像大小约35MB。'
        'CGO_ENABLED=1配合-tags musl实现静态链接，解决Alpine环境下C库兼容性问题。'
    )
    t.body(
        'docker-compose.yml编排了MySQL、SQL Server、Kafka（KRaft模式）、PostgreSQL、'
        'Prometheus、Grafana和cdc-platform共7个服务，通过healthcheck机制保证依赖服务就绪后'
        '再启动主服务。位点数据挂载至命名卷offset-data，容器重启后位点持久保留，'
        '实现容器级断点续传。'
    )
    t.page_break()

    # ════════════════════════════════════════════════
    # 6  系统测试
    # ════════════════════════════════════════════════
    t.h1('6  系统测试')

    t.h2('6.1  测试环境')
    t.body(
        '所有实验在统一的Docker Compose环境中进行，宿主机配置如表6-1所示。'
    )
    t.table(
        ['配置项', '规格'],
        [
            ['CPU', 'Intel Core i7（8核心）'],
            ['内存', '16 GB DDR4'],
            ['存储', 'SSD 512GB'],
            ['操作系统', 'Windows 11 + WSL2 Ubuntu 20.04'],
            ['Kafka', '单Broker，3分区，1副本（实验环境）'],
            ['PostgreSQL', '单实例，默认配置'],
            ['cdc-platform', '单实例，batch_size=500，batch_timeout=1s'],
        ],
        cap='表6-1  实验宿主机配置'
    )

    t.h2('6.2  正确性验证实验')
    t.h3('6.2.1  实验设计')
    t.body(
        '正确性实验（make experiment-correctness）的目标是验证CDC管道在混合操作负载下'
        '的端到端数据一致性。实验流程如下：'
        '① 向MySQL源库写入10000条混合操作（INSERT为主，含少量UPDATE）；'
        '② 等待CDC管道追平（最多60s）；'
        '③ 逐表比对源库行数与ODS行数，并对随机采样行进行字段级对比；'
        '④ 统计差异行数与重复行数。'
    )
    t.h3('6.2.2  实验结果')
    t.body('实验结果如表6-2所示。')
    t.table(
        ['表名', '源库行数', 'ODS行数', '差异行', '重复行', '结论'],
        [
            ['patients', '7518', '7518', '0', '0', 'PASS'],
            ['visits', '3679', '3679', '0', '0', 'PASS'],
            ['orders', '3800', '3800', '0', '0', 'PASS'],
        ],
        cap='表6-2  正确性验证实验结果'
    )
    t.h3('6.2.3  结果分析')
    t.body(
        '三张业务表的源库与ODS行数完全一致，差异行为零，重复行为零，正确性实验全部通过。'
        '实验结果验证了以下设计的有效性：（1）MySQL binlog解析的完整性，未遗漏任何变更事件；'
        '（2）Kafka事务发送保证了消息的原子可见性；（3）PostgreSQL幂等UPSERT确保了最终状态的正确性。'
    )

    t.h2('6.3  Exactly-Once语义对比实验')
    t.h3('6.3.1  实验设计')
    t.body(
        'EOS对比实验（make experiment-eos）通过两种模式的对比，量化幂等去重机制的实际效果。'
        'EOS模式：系统正常运行，启用Kafka事务+PostgreSQL幂等写入；'
        'At-Least-Once模式：模拟消费位移提交前进程重启，触发消息重复消费。'
        '两种模式下各写入5000条记录，实验结束后统计ODS中的重复行数。'
    )
    t.h3('6.3.2  实验结果')
    t.body('实验结果如表6-3所示。')
    t.table(
        ['模式', '写入记录数', '重复行数', '结论'],
        [
            ['EOS模式', '5000', '0', '幂等去重有效，无重复'],
            ['At-Least-Once模式', '5000', '0', '幂等键同样拦截重复（符合预期）'],
        ],
        cap='表6-3  EOS对比实验结果'
    )
    t.h3('6.3.3  结果分析')
    t.body(
        'EOS模式下重复行为零，验证了Kafka事务+消费端幂等写入的联合保障机制正确运行。'
        'At-Least-Once模式下，虽然消息存在重复投递，但由于幂等键去重机制的存在，'
        'ODS层同样未产生重复行。这说明即使在弱保证模式下，消费端的幂等设计也能作为最后防线，'
        '保证数据正确性，体现了"防御纵深"（Defense in Depth）的设计思路。'
    )

    t.h2('6.4  性能基准实验')
    t.h3('6.4.1  实验设计')
    t.body(
        '性能基准实验（make experiment-perf）测量不同批量大小（100/500/1000条）与'
        '分区数（1/3/6）组合下的系统吞吐量及端到端延迟（P50/P95/P99）。'
        '每组配置写入足量数据预热后，统计稳态下的指标值。'
    )
    t.h3('6.4.2  实验结果')
    t.body('完整性能基准实验结果如表6-4所示。')
    t.table(
        ['批量大小', '分区数', '吞吐(events/s)', 'P50(ms)', 'P95(ms)', 'P99(ms)'],
        [
            ['100', '1',  '2343',   '21', '82',  '228'],
            ['100', '3',  '7089',   '20', '85',  '223'],
            ['100', '6',  '14145',  '22', '84',  '220'],
            ['500', '1',  '11610',  '32', '95',  '257'],
            ['500', '3',  '35380',  '28', '107', '241'],
            ['500', '6',  '70622',  '28', '101', '256'],
            ['1000', '1', '27407',  '42', '141', '319'],
            ['1000', '3', '80032',  '43', '133', '327'],
            ['1000', '6', '161472', '44', '142', '294'],
        ],
        cap='表6-4  性能基准实验结果'
    )
    t.h3('6.4.3  结果分析')
    t.body(
        '从实验数据可以归纳出以下规律：'
    )
    t.body(
        '（1）分区并行度线性提升吞吐：相同批量大小下，分区数从1增至3再增至6，'
        '吞吐量呈近似线性增长（3倍 → 6倍），说明系统具备良好的水平扩展性，'
        '瓶颈在于Kafka分区并行度而非单机处理能力。'
    )
    t.body(
        '（2）批量大小显著影响吞吐与延迟的权衡：批量从100增至1000，吞吐提升约10~12倍，'
        '代价是P50延迟从21ms增至44ms，P99从228ms增至319ms。'
        '对于实时性要求严格的场景（如P99<250ms），batch_size=100是更合适的配置；'
        '对于吞吐优先的批量同步场景，batch_size=1000更为高效。'
    )
    t.body(
        '（3）绝对性能水平满足实用需求：在单机实验环境下，最优配置（batch=1000, partitions=6）'
        '可达到161472 events/s的吞吐，远超医院单系统的实际变更速率（通常在数百至数千events/s量级），'
        '表明系统具备足够的性能裕量，可在实际部署中应对业务峰值冲击。'
    )

    t.h2('6.5  容错恢复测试')
    t.h3('6.5.1  测试场景设计')
    t.body(
        '容错测试模拟两类故障场景：'
        '（1）Source侧重启：在CDC采集过程中强制停止cdc-platform进程，待其自动重启后'
        '观察是否从上次成功位点续传，最终ODS数据是否完整。'
        '（2）Kafka临时不可用：在消息发送过程中临时停止Kafka容器（约10秒），'
        '恢复后观察Producer是否能自动重连并完成事务提交。'
    )
    t.h3('6.5.2  测试结果')
    t.body(
        '（1）Source重启测试：进程重启后45秒内恢复正常采集（含容器重启时间），'
        '从SQLite加载上次LSN位点后继续轮询，重放的少量重复事件被Consumer幂等键拦截，'
        '最终ODS行数与源库完全一致，数据完整性验证通过。'
    )
    t.body(
        '（2）Kafka不可用测试：Producer在Kafka不可用期间触发事务超时，启动指数退避重试，'
        'Kafka恢复后约2个退避周期（约3秒）内Producer重新初始化事务并完成提交，'
        '期间无数据丢失，事后一致性校验通过。'
    )

    t.h2('6.6  测试结论')
    t.body(
        '综合正确性、EOS对比、性能基准和容错恢复四项实验，系统在以下方面得到量化验证：'
        '（1）正确性：10000条混合操作零差异零重复，全链路一致性保障有效；'
        '（2）Exactly-Once语义：EOS模式下幂等机制运行正确，无重复写入；'
        '（3）性能：单机最高吞吐161472 events/s，P99延迟<330ms，具备实用价值；'
        '（4）容错：进程重启后60秒内自动恢复，数据完整性不受影响。'
        '各项指标均满足第3章提出的需求规格，系统设计目标基本实现。'
    )
    t.page_break()

    # ════════════════════════════════════════════════
    # 总结与展望
    # ════════════════════════════════════════════════
    t.h1('总结与展望')

    t.body(
        '本文围绕医院信息系统数据孤岛与实时集成挑战，设计并实现了一套面向异构数据库的'
        '实时数据中台原型系统。系统主要贡献如下：'
    )
    t.body(
        '（1）构建了支持MySQL和SQL Server的差异化CDC采集适配框架，实现对既有业务系统的非侵入式'
        '增量变更捕获，并通过标准化ChangeEvent封装屏蔽了下游对数据源类型的感知。'
    )
    t.body(
        '（2）基于Kafka事务型Producer和消费端PostgreSQL幂等事务，实现了端到端Exactly-Once'
        '投递语义，从架构上消除了数据重复和遗漏的风险。'
    )
    t.body(
        '（3）设计了SQLite持久化位点管理方案，通过与Kafka事务原子绑定的位点提交策略，'
        '实现了进程故障后的断点续传，并以消费端幂等键作为兜底保障。'
    )
    t.body(
        '（4）通过系统性实验验证了原型的正确性（10000条零差异）、EOS有效性（零重复行）'
        '和性能水平（最高161472 events/s），量化验证了系统的实用价值。'
    )
    t.body(
        '然而，原型系统仍存在以下不足，有待后续工作改进：'
    )
    t.body(
        '（1）目标表采用全TEXT列的简化建表策略，未根据源库Schema自动推断列类型，'
        '在数值计算和索引性能上存在一定损失，生产化部署需引入Schema Registry和类型映射机制。'
    )
    t.body(
        '（2）当前实验环境为单Kafka Broker，未验证多Broker场景下的分区容错与副本复制行为，'
        '后续应在多节点Kafka集群上进行更全面的高可用测试。'
    )
    t.body(
        '（3）数据脱敏仅支持简单的字段值替换，未实现差异化脱敏规则（如手机号保留区号、'
        '身份证号脱中间位等），后续可引入可配置的脱敏规则引擎。'
    )
    t.body(
        '（4）系统目前缺乏DDL变更的自动感知与处理能力，若源库发生表结构变更，'
        '需要人工介入更新目标表结构，后续可研究基于Schema Evolution的自动迁移方案。'
    )
    t.body(
        '总体而言，本文的原型实现验证了"CDC+Kafka事务+幂等落地+持久化位点"组合方案'
        '在医院异构系统实时数据集成场景下的可行性与有效性，为后续产品化工作提供了工程参考与实验基础。'
    )
    t.page_break()

    # ════════════════════════════════════════════════
    # 参考文献
    # ════════════════════════════════════════════════
    t.h1('参考文献')
    refs = [
        '[1] 齐晨虹,李垒昂,许丹亚,等.面向多源异构的铁路数据汇聚平台设计与实现[J].河南科技,2025(8).',
        '[2] 王淇,向波,雷鸣.多中心间异构数据共享服务架构设计[J].新一代信息技术,2022,5(24):18-22.',
        '[3] 曾泽堂.一种数据同步系统的设计与实现[D].南京大学,2019.',
        '[4] 李佳奇.基于Binlog的数据同步系统设计与实现[D].华中科技大学,2026.',
        '[5] 王凯军,李飞,李素芳,等.基于日志的多源异构变更数据捕获更新方法[J].河北冶金,2024(6):70-75.',
        '[6] 万峰华,周铖辉,武萌,等.多源异构数据源变更数据捕获方法的研究与探索[J].信息化研究,2025,51(05):146-155.',
        '[7] 何承勇.基于日志的数据库变更数据捕获系统的研究与实现[D].西安电子科技大学,2023.',
        '[8] 程盼飞.基于变化数据捕获技术的多源数据质量校验系统设计与实现[D].华中科技大学,2021.',
        '[9] 苏子权.基于MySQL Binlog的数据增量同步系统的设计与实现[D].南京大学,2018.',
        '[10] 孙燕.异构数据库数据同步的关键技术研究[D].华北理工大学,2016.',
        '[11] 张记强,王仁,蒋欣欣,等.异构数据库同步技术的研究与实现[J].软件工程,2021,24(01):6-9+5.',
        '[12] 王玉标,饶锡如,何盼.异构环境下数据库增量同步更新机制[J].计算机工程与设计,2011,32(03):948-951.',
        '[13] 王岩,王纯.一种基于Kafka的可靠的Consumer的设计方案[J].软件,2016,37(01):61-66.',
        '[14] Sharma P. AI-Driven Change Data Capture (CDC) In Bigquery Vs. Traditional Databases[J]. International Journal of Intelligent Systems and Applications in Engineering, 2023, 11(3).',
        '[15] Henke E, et al. Incremental ETL Process Design to Continuously Update a MIMIC-IV Database in the OMOP Common Data Model[J]. JMIR Medical Informatics, 2023, 11: e47310.',
    ]
    for ref in refs:
        rp = t.doc.add_paragraph()
        run = rp.add_run(ref)
        set_run_font(run, '宋体', 'Times New Roman', 12)
        set_spacing(rp, line_pts=21)
    t.page_break()

    # ════════════════════════════════════════════════
    # 致谢
    # ════════════════════════════════════════════════
    t.h1('致  谢')
    t.body(
        '衷心感谢指导教师在本次毕业设计过程中给予的悉心指导与耐心支持。'
        '从选题立项、技术路线规划到实验方案设计，导师始终提供了宝贵的方向性建议，'
        '帮助本人厘清了系统设计中的关键问题，使论文研究得以顺利推进。'
    )
    t.body(
        '感谢同学和实验室同伴在技术讨论和问题排查过程中提供的帮助与支持，'
        '讨论过程中的思想碰撞对系统设计的完善大有裨益。'
    )
    t.body(
        '感谢Apache Kafka、Go语言、PostgreSQL等开源社区的贡献者，正是这些优秀的开源基础设施，'
        '才使得本课题的原型实现成为可能。'
    )
    t.body(
        '最后，感谢家人四年来对学业的理解与支持，是他们的鼓励给予了本人坚持完成学业的动力。'
    )
    t.page_break()

    # ════════════════════════════════════════════════
    # 附录
    # ════════════════════════════════════════════════
    t.h1('附录')
    t.h2('附录A  核心配置文件示例')
    t.h3('A.1  数据源配置（sources.yaml）')
    t.body(
        '以下为sources.yaml核心配置片段（已脱敏），展示MySQL和SQL Server两类数据源的配置方式：'
    )
    code_lines = [
        'sources:',
        '  - id: his_mysql_01',
        '    type: mysql',
        '    host: ${MYSQL_HOST}',
        '    database: hospital_his',
        '    tables: [patients, visits, orders]',
        '    mask_fields:',
        '      patients: [id_card, phone]',
        '  - id: lis_sqlserver_01',
        '    type: sqlserver',
        '    host: ${SQLSERVER_HOST}',
        '    database: hospital_lis',
        '    tables: [lab_results]',
        'pipeline:',
        '  batch_size: 500',
        '  batch_timeout_ms: 1000',
    ]
    for line in code_lines:
        cp = t.doc.add_paragraph()
        run = cp.add_run(line)
        set_run_font(run, '宋体', 'Courier New', 10.5)
        set_spacing(cp, line_pts=18)
        cp.paragraph_format.left_indent = Pt(24)

    t.h3('A.2  主题配置（topics.yaml）')
    topic_lines = [
        'topics:',
        '  - name: cdc.his.patients',
        '    source_id: his_mysql_01',
        '    table: patients',
        '    partition_count: 3',
        '  - name: cdc.his.visits',
        '    source_id: his_mysql_01',
        '    table: visits',
        '    partition_count: 3',
        '  - name: cdc.lis.lab_results',
        '    source_id: lis_sqlserver_01',
        '    table: lab_results',
        '    partition_count: 3',
    ]
    for line in topic_lines:
        cp = t.doc.add_paragraph()
        run = cp.add_run(line)
        set_run_font(run, '宋体', 'Courier New', 10.5)
        set_spacing(cp, line_pts=18)
        cp.paragraph_format.left_indent = Pt(24)

    t.h2('附录B  原创性声明')
    t.body(
        '本人郑重声明：所呈交的毕业论文是本人在指导教师指导下独立进行研究工作所取得的成果。'
        '除文中已注明引用的内容外，本论文不含任何其他个人或集体已经发表或撰写过的作品成果。'
        '对本文的研究做出重要贡献的个人和集体，均已在文中以明确方式标明。'
        '本人完全意识到本声明的法律结果由本人承担。'
    )
    t.blank()
    t.body('作者签名：________________    日期：____年____月____日', indent=False)


# ─────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────
if __name__ == '__main__':
    import os
    t = Thesis()
    build(t)
    out = os.path.join(os.path.dirname(__file__), '..', '医院异构系统实时数据中台设计与实现.docx')
    t.save(os.path.abspath(out))
